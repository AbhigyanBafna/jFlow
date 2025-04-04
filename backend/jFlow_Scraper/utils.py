import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches

def clean_text(text):
    # Remove HTML tags
    text = re.sub(r'<[^>]*?>', '', text)
    # Remove URLs
    text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
    # Remove special characters
    text = re.sub(r'[^a-zA-Z0-9 ]', '', text)
    # Replace multiple spaces with a single space
    text = re.sub(r'\s{2,}', ' ', text)
    # Trim leading and trailing whitespace
    text = text.strip()
    return text

def create_formatted_cover_letter_docx(cover_letter_text, filename='cover_letter.docx'):
    # Create a new Document
    doc = Document()
    
    # Adjust left and right margins
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.5)  # Set left margin to 0.5 inches
        section.right_margin = Inches(0.5)  # Set right margin to 0.5 inches

    # Split the cover letter into paragraphs
    cover_letter_sections = cover_letter_text.split("\n\n")
    length = len(cover_letter_sections)

    # Check if the sections exist (address, recipient, body) and add them
    if length >= 3:
        # First part: Header with contact information, right-aligned, user's contact details
        header = doc.add_paragraph()
        header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        first_bold = False
        for line in cover_letter_sections[0].split("\n"):
            header.add_run(f"{line}\n").bold = not first_bold
            first_bold = True
        header.style.font.size = Pt(14)
        header.style.font.name = "Calibri"

        # Second part: Recipient's contact information, left-aligned
        recipient = doc.add_paragraph()
        recipient.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        for line in cover_letter_sections[1].split("\n"):
            recipient.add_run(f"{line}\n").bold
            first_bold = True

        recipient.style.font.size = Pt(14)
        
        # Third part: salutation
        salutation = doc.add_paragraph(f"{cover_letter_sections[2]}\n")
        salutation.style.font.size = Pt(12)  # Set a default font size for the body

        # Third part to end-1: Body of the cover letter
        for section in cover_letter_sections[3:length-1]:
            # Instead of splitting by newlines, we replace them with spaces
            paragraph_text = section.replace("\n", ' ')

            # Add the whole section as one justified paragraph
            paragraph = doc.add_paragraph(paragraph_text)
            paragraph.style.font.size = Pt(12)  # Set font size for the body
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Justify the entire paragraph

        # End greeting
        greetings = doc.add_paragraph(f"{cover_letter_sections[length-1]}\n")
        greetings.style.font.size = Pt(12)  # Set a default font size for the body

    else:
        # If the format doesn't match expected sections, just add all text as body
        paragraph = doc.add_paragraph(cover_letter_text)
        paragraph.style.font.size = Pt(12)

    # Save the document
    doc.save(filename)
    return filename
